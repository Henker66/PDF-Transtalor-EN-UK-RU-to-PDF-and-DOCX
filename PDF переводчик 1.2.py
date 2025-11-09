
#SBA System
#09.11.2025
# PDF Translator EN-UK-RU
# Version 1.2

import fitz
from pdf2image import convert_from_path
import pytesseract
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from tkinter import filedialog, messagebox
import threading
from datetime import datetime
from PIL import Image, ImageDraw, ImageFont
import os
from deep_translator import GoogleTranslator
from docx import Document
import time
import gc
import shutil  # для очистки кеша

# Настройка Tesseract
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
FONT_PATH = r"C:\Windows\Fonts\arial.ttf"  # шрифт с поддержкой кириллицы
CACHE_DIR = os.path.join(os.getcwd(), "cache_translations")
os.makedirs(CACHE_DIR, exist_ok=True)

class PDFTranslatorApp:
    def __init__(self):
        self.root = ttk.Window(themename="cosmo")
        self.root.title("PDF → PDF / DOCX (Русский)")
        self.root.geometry("800x600")

        self.pdf_path = ""
        self.stop_flag = False
        self.translator = GoogleTranslator(source='auto', target='ru')

        self.create_widgets()
        self.root.mainloop()

    def create_widgets(self):
        ttk.Label(self.root, text="Выберите PDF файл:").pack(pady=10, anchor=W)
        frame = ttk.Frame(self.root)
        frame.pack(padx=10, fill="x")
        self.entry_pdf = ttk.Entry(frame, width=60)
        self.entry_pdf.pack(side="left", expand=True, fill="x")
        ttk.Button(frame, text="Обзор", command=self.browse_pdf).pack(side="left", padx=5)

        ttk.Button(self.root, text="Перевести PDF на русский", command=self.start_translation).pack(pady=5)
        ttk.Button(self.root, text="Экспорт в DOCX", command=self.start_docx_export).pack(pady=5)
        ttk.Button(self.root, text="Остановить", command=self.stop_translation).pack(pady=5)
        ttk.Button(self.root, text="Очистить кеш", command=self.clear_cache).pack(pady=5)

        ttk.Label(self.root, text="Прогресс:").pack(anchor=W, padx=10)
        frame_progress = ttk.Frame(self.root)
        frame_progress.pack(pady=5, fill="x")
        self.progress = ttk.Progressbar(frame_progress, length=700, mode="determinate")
        self.progress.pack(side="left")
        self.percent_label = ttk.Label(frame_progress, text="0%")
        self.percent_label.pack(side="left", padx=5)
        self.time_label = ttk.Label(frame_progress, text="Оставшееся время: --:--")
        self.time_label.pack(side="left", padx=10)

        ttk.Label(self.root, text="Лог:").pack(anchor=W, padx=10)
        self.log_box = ttk.ScrolledText(self.root, height=20)
        self.log_box.pack(fill="both", expand=True, padx=10)

    def log(self, message):
        timestamp = datetime.now().strftime("%H:%M:%S")
        self.log_box.insert("end", f"[{timestamp}] {message}\n")
        self.log_box.see("end")

    def browse_pdf(self):
        path = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf")])
        if path:
            self.pdf_path = path
            self.entry_pdf.delete(0, "end")
            self.entry_pdf.insert(0, path)

    def stop_translation(self):
        self.stop_flag = True
        self.log("⛔ Остановка...")

    def clear_cache(self):
        if os.path.exists(CACHE_DIR):
            shutil.rmtree(CACHE_DIR)
        os.makedirs(CACHE_DIR, exist_ok=True)
        self.log("🗑 Кеш очищен")
        messagebox.showinfo("Готово", "Кеш очищен!")

    # ---------------- PDF перевод ----------------
    def start_translation(self):
        if not self.pdf_path:
            messagebox.showwarning("Ошибка", "Выберите PDF файл")
            return
        self.stop_flag = False
        threading.Thread(target=self.translate_pdf, daemon=True).start()

    def draw_text_on_image(self, img, text, font_path=FONT_PATH, font_size=24, margin=50):
        draw = ImageDraw.Draw(img)
        font = ImageFont.truetype(font_path, size=font_size)

        wrapped_lines = []
        for paragraph in text.split("\n"):
            paragraph = paragraph.strip()
            if not paragraph:
                wrapped_lines.append("")
                continue
            line = paragraph
            while draw.textlength(line, font=font) > img.width - 2*margin:
                for cut in reversed(range(1, len(line))):
                    if line[cut] == " ":
                        wrapped_lines.append(line[:cut])
                        line = line[cut+1:]
                        break
                else:
                    wrapped_lines.append(line[:len(line)//2])
                    line = line[len(line)//2:]
            wrapped_lines.append(line)
            wrapped_lines.append("")

        y = margin
        line_spacing = int(font_size*0.2)
        for line in wrapped_lines:
            if y > img.height - margin:
                break
            draw.text((margin, y), line, font=font, fill=(0,0,0))
            bbox = draw.textbbox((0,0), line, font=font)
            line_height = bbox[3]-bbox[1]+line_spacing
            y += line_height

        return img

    def translate_pdf(self):
        # Автоочистка кеша перед новым переводом
        self.clear_cache()

        input_pdf = self.pdf_path
        output_pdf = input_pdf.replace(".pdf", "_translated.pdf")
        self.log(f"🔄 Начало перевода PDF: {input_pdf}")

        try:
            images = convert_from_path(input_pdf)
        except Exception as e:
            self.log(f"[Ошибка] Не удалось конвертировать PDF: {e}")
            return

        total_pages = len(images)
        self.progress["maximum"] = total_pages
        translated_images = []
        start_time = time.time()

        for i, img in enumerate(images):
            if self.stop_flag:
                self.log("⛔ Перевод остановлен")
                break

            self.log(f"Обрабатываем страницу {i+1}/{total_pages}")

            text = pytesseract.image_to_string(img, lang="eng+ukr+rus")
            try:
                translated_text = self.translator.translate(text)
            except Exception as e:
                self.log(f"[Ошибка перевода страницы {i+1}] {e}")
                translated_text = text

            overlay = Image.new("RGB", img.size, (255,255,255))
            overlay = self.draw_text_on_image(overlay, translated_text)
            translated_images.append(overlay)

            self.progress["value"] = i+1
            percent = int(((i+1)/total_pages)*100)
            self.percent_label.config(text=f"{percent}%")

            elapsed = time.time() - start_time
            avg_per_page = elapsed / (i+1)
            remaining = int(avg_per_page * (total_pages - (i+1)))
            self.time_label.config(text=f"Оставшееся время: {remaining//60:02d}:{remaining%60:02d}")

        # Сохраняем PDF
        try:
            if translated_images:
                translated_images[0].save(
                    output_pdf,
                    save_all=True,
                    append_images=translated_images[1:]
                )
                self.log(f"✅ PDF успешно переведён: {output_pdf}")
                messagebox.showinfo("Готово", f"PDF переведён!\nСохранён в:\n{output_pdf}")
        except Exception as e:
            self.log(f"[Ошибка сохранения PDF] {e}")
        finally:
            # Очистка памяти
            for im in images:
                im.close()
            for im in translated_images:
                im.close()
            images.clear()
            translated_images.clear()
            del images, translated_images
            gc.collect()

    # ---------------- DOCX экспорт ----------------
    def start_docx_export(self):
        if not self.pdf_path:
            messagebox.showwarning("Ошибка", "Выберите PDF файл")
            return
        self.stop_flag = False
        threading.Thread(target=self.export_to_docx, daemon=True).start()

    def export_to_docx(self):
        # Автоочистка кеша перед экспортом
        self.clear_cache()

        input_pdf = self.pdf_path
        output_docx = input_pdf.replace(".pdf", "_translated.docx")
        self.log(f"🔄 Начало экспорта в DOCX: {input_pdf}")

        try:
            images = convert_from_path(input_pdf)
        except Exception as e:
            self.log(f"[Ошибка] Не удалось конвертировать PDF: {e}")
            return

        total_pages = len(images)
        self.progress["maximum"] = total_pages
        doc = Document()
        start_time = time.time()

        for i, img in enumerate(images):
            if self.stop_flag:
                self.log("⛔ Экспорт остановлен")
                break

            self.log(f"Обрабатываем страницу {i+1}/{total_pages}")

            text = pytesseract.image_to_string(img, lang="eng+ukr+rus")
            try:
                translated_text = self.translator.translate(text)
            except Exception as e:
                self.log(f"[Ошибка перевода страницы {i+1}] {e}")
                translated_text = text

            for paragraph in translated_text.split("\n"):
                paragraph = paragraph.strip()
                if paragraph:
                    doc.add_paragraph(paragraph)
            doc.add_paragraph("")  # разрыв страниц

            self.progress["value"] = i+1
            percent = int(((i+1)/total_pages)*100)
            self.percent_label.config(text=f"{percent}%")

            elapsed = time.time() - start_time
            avg_per_page = elapsed / (i+1)
            remaining = int(avg_per_page * (total_pages - (i+1)))
            self.time_label.config(text=f"Оставшееся время: {remaining//60:02d}:{remaining%60:02d}")

        try:
            doc.save(output_docx)
            self.log(f"✅ DOCX успешно сохранён: {output_docx}")
            messagebox.showinfo("Готово", f"DOCX сохранён!\nСохранён в:\n{output_docx}")
        except Exception as e:
            self.log(f"[Ошибка сохранения DOCX] {e}")
        finally:
            for im in images:
                im.close()
            images.clear()
            del images
            gc.collect()

if __name__ == "__main__":
    PDFTranslatorApp()
